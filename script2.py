from docx.shared import Pt  ,RGBColor  
from docx import Document, document
import openpyxl
import os
from txtparse import parse_txt_to_data
mf=["LUMIÉRE TRANSPORT - TUNIS 616717/D/A/M/000","LOGISTIQUE MAGASIN GENERAL 1480261/Y/A/M/000"]
# paths?
template_path = "./ChangeMe.docx"
output_dir = "factures"
os.makedirs(output_dir, exist_ok=True)
excel_file = "word_metadata.xlsx"

if os.path.exists(excel_file):
    wb = openpyxl.load_workbook(excel_file)
    ws = wb.active
else:
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.append(["File Name", "Content Summary"])

#user inputs?
data = parse_txt_to_data("data.txt")
num_facture = "00390"
date_facture = "14/12/2024"
Marque = "Camion"
Matricule = "1234 tu 489"
client="lum"



doc = Document(template_path)
table = doc.tables[1] 
cell=doc.tables[0].cell(1,0)
cell2=doc.tables[0].cell(2,0)
cell3=doc.tables[0].cell(2,2)

para = cell.paragraphs[0]
para2 = cell2.paragraphs[4]
para3 = cell3.paragraphs[0]
para3.add_run(client).font.size = Pt(9)

text = f"NUM: {num_facture}\nDATE: {date_facture}"
text2=f"Marque: {Marque}\nMatricule: {Matricule}"

lines = text.split("\n")
lines2 = text2.split("\n")


para2.clear()
cell2.paragraphs[5].clear()
label_marque, value_marque = lines2[0].split(":")
run_marque_label = para2.add_run(label_marque+" : ")
run_marque_label.bold = True
run_marque_label.font.name = 'Century Gothic'
run_marque_label.font.color.rgb = RGBColor(0, 0, 0) 
run_marque_label.font.size = Pt(12) 
para2.add_run(Marque).font.size = Pt(12)

para2.add_run("\n")

label_mat, value_mat = lines2[1].split(":")
run_mat_label = para2.add_run(label_mat+" : ")
run_mat_label.bold = True
run_mat_label.font.name = 'Century Gothic'
run_mat_label.font.color.rgb = RGBColor(0, 0, 0) 
run_mat_label.font.size = Pt(12) 
para2.add_run(Matricule).font.size = Pt(12)


run_num_label = para.add_run(lines[0])
run_num_label.bold = True
run_num_label.font.name = 'Century Gothic'
run_num_label.font.color.rgb = RGBColor(0, 0, 0) 
run_num_label.font.size = Pt(12) 

para.add_run("\n")

run_date_label = para.add_run(lines[1])
run_date_label.bold = True
run_date_label.font.name = 'Century Gothic'
run_date_label.font.color.rgb = RGBColor(0, 0, 0) # Black
run_date_label.font.size = Pt(9)
j=1
for factures in data:
    i=1
    for facture in factures:
        values = list(facture.values())
        for col_index, value in enumerate(values):
            table.cell(i, col_index+1).text = value
        i=i+1
    file_name = f"Facture{j}.docx"
    j+=1
    file_path = os.path.join(output_dir, file_name)
    doc.save(file_path)




print("Documents created successfully.")

